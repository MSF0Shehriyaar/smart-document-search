import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ---- Page Setup ----
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# Set 2 columns for the first section
sectPr = doc.sections[0]._sectPr
cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="708"/>')
sectPr.append(cols)

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

def add_heading_custom(text, level=1, bold=True, italic=False, center=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    if size:
        run.font.size = Pt(size)
    elif level == 1:
        run.font.size = Pt(11)
    elif level == 2:
        run.font.size = Pt(10)
    elif level == 3:
        run.font.size = Pt(10)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_para(text, bold=False, italic=False, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_mixed_para(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2, space_before=0, size=10):
    """parts = list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_horizontal_rule():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/></w:pBdr>')
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

# ==================== TITLE BLOCK ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
run = p.add_run("Smart Document Search: An AI-Ready Information Retrieval System with TF-IDF Ranking and Modern Web Architecture")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Mohammed Shehriyaar F, Mohammed Muteeb Ahmed, Mohammed Ibrahim")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run("Department of Information Science")
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

add_horizontal_rule()

# ==================== ABSTRACT ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Abstract")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

add_para(
    "This paper presents the design and implementation of Smart Document Search, "
    "a full-stack web application that enables users to upload PDF and text documents "
    "and perform intelligent search queries using the Term Frequency-Inverse Document "
    "Frequency (TF-IDF) algorithm. The system features a modern cyberpunk-inspired "
    "glassmorphism user interface built with React 19 and Tailwind CSS, an Express.js "
    "backend with in-memory document storage, and real-time relevance scoring with "
    "animated result visualization. The architecture is designed to be AI-ready, "
    "incorporating the Google GenAI SDK for future semantic search enhancements. "
    "This paper discusses the system architecture, the TF-IDF information retrieval "
    "methodology, implementation details, and potential future directions for "
    "AI-powered document analysis.",
    size=9
)

add_mixed_para([
    ("Keywords: ", True, False),
    ("Information Retrieval, TF-IDF, Document Search, React, Express, PDF Parsing, Full-Text Search, Vector Space Model", False, True)
], size=9)

add_horizontal_rule()

# ==================== 1. INTRODUCTION ====================
add_heading_custom("1. Introduction", level=1, center=True, size=11)

add_heading_custom("1.1 Background", level=2, italic=True)
add_para(
    "In the modern digital era, the volume of unstructured text data generated by "
    "individuals and organizations has grown exponentially. Documents such as PDFs, "
    "text files, and reports contain valuable information, but extracting relevant "
    "content efficiently remains a significant challenge. Traditional file-system "
    "searches rely on filename matching or basic string containment, which often "
    "fail to capture semantic relevance or rank results by importance."
)
add_para(
    "Information Retrieval (IR) systems address this challenge by employing algorithms "
    "that measure the relevance between a user's query and the contents of a document "
    "collection. Among the foundational techniques in IR, the TF-IDF weighting scheme "
    "remains widely used due to its simplicity, interpretability, and effectiveness "
    "for many practical applications."
)

add_heading_custom("1.2 Problem Statement", level=2, italic=True)
add_para(
    "Existing document search solutions often fall into two categories: (1) overly "
    "simplistic tools that lack ranking capabilities, or (2) complex enterprise search "
    "platforms that require significant infrastructure (databases, search engines like "
    "Elasticsearch, cloud services). There is a need for a lightweight, self-contained "
    "document search system that can run locally, process common document formats, and "
    "provide ranked search results without external dependencies."
)

add_heading_custom("1.3 Objectives", level=2, italic=True)
add_para("The primary objectives of this project are:")
objectives = [
    "To develop a full-stack web application for uploading and searching PDF and text documents.",
    "To implement a TF-IDF-based ranking algorithm for relevance scoring.",
    "To design an intuitive, visually engaging user interface with real-time feedback.",
    "To architect the system for future integration with Large Language Models (LLMs) and semantic search."
]
for i, obj in enumerate(objectives, 1):
    add_para(f"{i}. {obj}", space_before=1)

add_heading_custom("1.4 Scope", level=2, italic=True)
add_para("The current implementation supports:")
scope = [
    "Multi-file upload of PDF and TXT documents",
    "In-memory storage and indexing",
    "Real-time TF-IDF search with relevance ranking",
    "Responsive glassmorphism UI with motion animations",
    "AI SDK integration readiness for future enhancements"
]
for item in scope:
    add_para(f"\u2022 {item}", space_before=1)

# ==================== 2. LITERATURE REVIEW ====================
add_heading_custom("2. Literature Review", level=1, center=True, size=11)

add_heading_custom("2.1 Information Retrieval and the Vector Space Model", level=2, italic=True)
add_para(
    "Information Retrieval (IR) is the science of searching for information in large "
    "collections of unstructured data [1]. The Vector Space Model (VSM), introduced by "
    "Gerard Salton in the 1970s, represents documents and queries as vectors in a "
    "high-dimensional space where each dimension corresponds to a term in the corpus [2]. "
    "The relevance between a query and a document is computed as a function of their "
    "vector representations."
)

add_heading_custom("2.2 TF-IDF Weighting", level=2, italic=True)
add_para(
    "Term Frequency-Inverse Document Frequency (TF-IDF) is a numerical statistic that "
    "reflects how important a word is to a document in a collection [3]. It is the "
    "product of two metrics:"
)
add_mixed_para([
    ("Term Frequency (TF): ", True, False),
    ("Measures how frequently a term appears in a document.", False, False)
], space_before=2)
add_mixed_para([
    ("Inverse Document Frequency (IDF): ", True, False),
    ("Measures how rare a term is across the entire corpus.", False, False)
], space_before=2)
add_para(
    "The TF-IDF weight increases proportionally to the number of times a word appears "
    "in the document but is offset by the frequency of the word in the corpus, which "
    "helps to control for the fact that some words are generally more common than others."
)

add_heading_custom("2.3 Modern Web Architectures for IR Applications", level=2, italic=True)
add_para(
    "Contemporary web applications for document processing typically employ a "
    "client-server architecture with a Single Page Application (SPA) frontend and a "
    "RESTful API backend [4]. React, combined with build tools like Vite, has become "
    "a standard for frontend development due to its component-based architecture and "
    "fast development experience [5]. On the backend, Node.js with Express provides "
    "a lightweight, event-driven runtime suitable for I/O-bound operations such as "
    "file uploads and text processing."
)

add_heading_custom("2.4 PDF Text Extraction", level=2, italic=True)
add_para(
    "Extracting text from PDF documents is a non-trivial task due to the format's "
    "complex internal structure. Libraries such as pdf-parse (based on Mozilla's PDF.js) "
    "provide robust capabilities for rendering PDF content to plain text, enabling "
    "downstream text analysis and search operations [6]."
)

# ==================== 3. SYSTEM ARCHITECTURE ====================
add_heading_custom("3. System Architecture", level=1, center=True, size=11)

add_heading_custom("3.1 Overview", level=2, italic=True)
add_para(
    "Smart Document Search follows a three-tier architecture consisting of:"
)
add_para("1. Presentation Layer: React-based SPA with Tailwind CSS styling", space_before=1)
add_para("2. Application Layer: Express.js server handling HTTP requests, file uploads, and search logic", space_before=1)
add_para("3. Data Layer: In-memory document store with runtime indexing", space_before=1)

add_heading_custom("3.2 Technology Stack", level=2, italic=True)

# Table: Technology Stack
table = doc.add_table(rows=11, cols=3)
table.style = 'Table Grid'
headers = ['Component', 'Technology', 'Purpose']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'

data = [
    ['Frontend', 'React 19', 'UI component rendering'],
    ['Build Tool', 'Vite 6.2', 'Module bundling and dev server'],
    ['Styling', 'Tailwind CSS 4', 'Utility-first CSS framework'],
    ['Animation', 'Motion 12', 'Declarative animations'],
    ['Icons', 'Lucide React', 'SVG icon library'],
    ['Backend', 'Express 4', 'HTTP server and routing'],
    ['Runtime', 'TypeScript 5.8', 'Type-safe development'],
    ['File Uploads', 'Multer 2.1', 'Multipart form handling'],
    ['PDF Parsing', 'pdf-parse 2.4', 'PDF text extraction'],
    ['AI SDK', '@google/genai 1.29', 'Future AI integration'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

add_heading_custom("3.3 Architectural Diagram", level=2, italic=True)
add_para(
    "The system architecture consists of a React client communicating with an Express "
    "server via HTTP/REST (JSON and FormData). The Express server handles file uploads "
    "through Multer, parses PDFs using pdf-parse for text extraction, and stores documents "
    "in an in-memory store. The TF-IDF Engine processes queries by tokenizing input, "
    "computing term frequency and inverse document frequency scores, and ranking results "
    "by relevance. In development mode, Vite middleware serves the React SPA from the "
    "same port. In production, static files from the dist/ directory are served.",
    size=9
)

add_heading_custom("3.4 Development and Production Modes", level=2, italic=True)
add_para(
    "The application uses Vite's middleware mode during development, allowing the Express "
    "server to serve both API routes and the React application from a single port (3000). "
    "In production, the React frontend is built to static files and served by Express from "
    "the dist/ directory."
)

# ==================== 4. METHODOLOGY ====================
add_heading_custom("4. Methodology", level=1, center=True, size=11)

add_heading_custom("4.1 Document Ingestion Pipeline", level=2, italic=True)
add_para("The document ingestion process consists of the following steps:")
ingestion = [
    ("File Selection: ", "Users select one or more files through a drag-and-drop interface or file picker."),
    ("Upload: ", "Files are transmitted to the server via HTTP POST with multipart/form-data encoding."),
    ("Parsing: ", "The server inspects each file's MIME type and extension. PDF files are parsed using pdf-parse; text files are read directly as UTF-8 strings."),
    ("Storage: ", "Extracted content is stored in an in-memory array alongside metadata (unique ID, filename, MIME type)."),
    ("Library Update: ", "The client fetches the updated document list and renders it in the sidebar."),
]
for i, (label, desc) in enumerate(ingestion, 1):
    add_mixed_para([(f"{i}. {label}", True, False), (desc, False, False)], space_before=2)

add_heading_custom("4.2 TF-IDF Algorithm Implementation", level=2, italic=True)
add_para(
    "The search algorithm implemented in the application follows the standard TF-IDF "
    "formulation with the following specifics:"
)

add_heading_custom("4.2.1 Tokenization", level=3, italic=True)
add_para(
    "Text is tokenized using a regular expression that extracts word boundaries. "
    "The tokenizer converts text to lowercase and extracts sequences of alphanumeric "
    "characters, filtering out punctuation and whitespace. This normalization ensures "
    "case-insensitive matching across documents and queries."
)

add_heading_custom("4.2.2 Term Frequency (TF)", level=3, italic=True)
add_para(
    "For a given term t and document d, the Term Frequency is computed as the ratio "
    "of the frequency of term t in document d to the total number of tokens in "
    "document d. This normalized frequency prevents bias toward longer documents, "
    "ensuring that the score reflects the proportion of the term relative to the "
    "document length."
)

add_heading_custom("4.2.3 Inverse Document Frequency (IDF)", level=3, italic=True)
add_para(
    "For a term t across corpus D, the Inverse Document Frequency is computed using "
    "a smoothed logarithmic formula. The smoothing terms (+1) prevent division by zero "
    "and ensure that terms appearing in all documents still receive a non-zero weight. "
    "A constant of 1 is added to the logarithm result to ensure IDF is always positive, "
    "so that even very common terms contribute a small positive weight."
)

add_heading_custom("4.2.4 Scoring Function", level=3, italic=True)
add_para(
    "The total relevance score for a document given a query is the sum of TF * IDF "
    "for all unique query terms. Documents with a score greater than zero are returned, "
    "sorted in descending order by score. This approach effectively combines local "
    "significance (TF) with global rarity (IDF) to produce a meaningful relevance ranking."
)

add_heading_custom("4.2.5 Result Enrichment", level=3, italic=True)
add_para("Each result includes:")
enrichment = [
    "Document ID and filename for identification",
    "Computed relevance score (floating-point, 4 decimal places)",
    "A 200-character text snippet from the beginning of the document",
    "Match confidence percentage derived from the score"
]
for item in enrichment:
    add_para(f"\u2022 {item}", space_before=1)

add_heading_custom("4.3 User Interface Design", level=2, italic=True)
add_para("The UI follows a cyberpunk-inspired aesthetic with the following design principles:")
ui_principles = [
    ("Glassmorphism: ", "Frosted glass panels with backdrop-blur and semi-transparent backgrounds"),
    ("Dark Theme: ", "Slate and indigo color palette with high contrast text"),
    ("Motion Design: ", "Staggered entrance animations for search results using Framer Motion"),
    ("Real-time Feedback: ", "Animated progress bars showing match confidence, loading spinners during operations, and live status indicators"),
    ("Responsive Layout: ", "Flexbox-based layout that adapts from single-column (mobile) to two-column (desktop)"),
]
for label, desc in ui_principles:
    add_mixed_para([("\u2022 " + label, True, False), (desc, False, False)], space_before=1)

# ==================== 5. IMPLEMENTATION ====================
add_heading_custom("5. Implementation", level=1, center=True, size=11)

add_heading_custom("5.1 Backend Implementation", level=2, italic=True)
add_para(
    "The Express server is written in TypeScript and provides three primary API endpoints "
    "for document management and search operations."
)

add_mixed_para([
    ("POST /api/upload \u2014 ", True, False),
    ("Accepts an array of files via Multer's memory storage. Each file is processed based "
     "on its type: PDF files are parsed using PDFParse with error handling for invalid or "
     "corrupted files; text files are converted from buffer to UTF-8 string. Processed "
     "documents are pushed to the in-memory documents array with a randomly generated ID.", False, False)
], space_before=2)

add_mixed_para([
    ("GET /api/documents \u2014 ", True, False),
    ("Returns a sanitized list of documents containing only id, name, and type fields, "
     "protecting the full text content from unnecessary exposure.", False, False)
], space_before=2)

add_mixed_para([
    ("POST /api/search \u2014 ", True, False),
    ("Accepts a JSON body with a query string. The server validates the query, tokenizes "
     "it along with all document contents, computes TF-IDF scores for each document, "
     "filters results with score > 0, sorts by score descending, and returns enriched "
     "result objects.", False, False)
], space_before=2)

add_heading_custom("5.2 Frontend Implementation", level=2, italic=True)
add_para(
    "The React frontend is implemented as a single functional component with state "
    "management using React hooks. The component state includes the current search query, "
    "ranked search results, uploaded document library, and operation status flags for "
    "upload and search. The component is structured into three main regions: a header "
    "with application branding and system status, a sidebar with file upload zone and "
    "active document library, and a main content area with search input form and results "
    "display."
)

add_heading_custom("5.3 PDF Parsing Challenges", level=2, italic=True)
add_para("PDF parsing presents several challenges that the application addresses:")
pdf_challenges = [
    "Invalid PDFs are caught and reported with descriptive error messages",
    "Corrupted files are detected via PDF.js exception types (InvalidPDFException, FormatError, AbortException)",
    "Non-PDF files with .pdf extensions are rejected with clear feedback"
]
for item in pdf_challenges:
    add_para(f"\u2022 {item}", space_before=1)

add_heading_custom("5.4 Error Handling", level=2, italic=True)
add_para("The application implements comprehensive error handling at multiple layers:")
errors = [
    "Network errors: Fetch failures are logged to the console with descriptive messages",
    "Upload errors: Server-side parsing errors are propagated to the client UI",
    "Validation errors: Missing files or queries return appropriate HTTP status codes (400 Bad Request)",
    "Runtime errors: PDF parsing exceptions are caught and returned as JSON error responses"
]
for item in errors:
    add_para(f"\u2022 {item}", space_before=1)

# ==================== 6. RESULTS AND FEATURES ====================
add_heading_custom("6. Results and Features", level=1, center=True, size=11)

add_heading_custom("6.1 Functional Features", level=2, italic=True)
add_para("The implemented system provides the following capabilities:")
features = [
    "Multi-File Upload: Users can upload multiple PDF and TXT files simultaneously",
    "Document Library: A persistent sidebar displays all uploaded documents with file type icons",
    "Real-Time Search: Queries are processed instantly against the in-memory index",
    "Relevance Ranking: Results are ordered by computed TF-IDF score",
    "Match Visualization: Animated confidence bars and percentage indicators provide intuitive relevance feedback",
    "Responsive Design: The application is usable on desktop and tablet viewports"
]
for item in features:
    add_mixed_para([("\u2022 ", False, False), (item, False, False)], space_before=1)

add_heading_custom("6.2 Performance Characteristics", level=2, italic=True)

# Table: Performance
table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Table Grid'
perf_headers = ['Metric', 'Value']
for i, h in enumerate(perf_headers):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'

perf_data = [
    ['Upload Processing', '< 1 second for typical documents'],
    ['Search Latency', 'Sub-millisecond for moderate corpus sizes'],
    ['File Size Limit', '10 MB per upload (configurable)'],
    ['Memory Usage', 'Proportional to total extracted text size'],
]
for row_idx, row_data in enumerate(perf_data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table2.rows[row_idx].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

add_heading_custom("6.3 User Interface", level=2, italic=True)
add_para(
    "The interface features a sticky header with aurora background effects, a drag-and-drop "
    "upload zone with hover animations, a scrollable document library with truncated filenames, "
    "a prominent search bar with focus states and ring effects, animated result cards with "
    "staggered entrance transitions, and a system footer displaying operational status."
)

# ==================== 7. FUTURE WORK ====================
add_heading_custom("7. Future Work", level=1, center=True, size=11)

add_heading_custom("7.1 Semantic Search with Gemini Embeddings", level=2, italic=True)
add_para(
    "The application includes the @google/genai SDK and Vite environment configuration for "
    "GEMINI_API_KEY. Future enhancements include embedding generation using Gemini's "
    "text-embedding models to convert documents and queries into dense vector representations, "
    "replacing TF-IDF with cosine similarity for semantic matching beyond lexical overlap, "
    "and hybrid scoring combining TF-IDF and embedding scores for optimal relevance."
)

add_heading_custom("7.2 AI-Generated Summaries", level=2, italic=True)
add_para(
    "Integration with Gemini's generative models would enable automatic document summarization, "
    "direct answers to user questions with citation to source documents, and AI-suggested "
    "query reformulations to improve recall."
)

add_heading_custom("7.3 Persistent Storage", level=2, italic=True)
add_para(
    "Current in-memory storage is ephemeral. Production deployment would require database "
    "integration (PostgreSQL or MongoDB) for document persistence, a vector database "
    "(Pinecone, Weaviate, or pgvector) for embedding storage and approximate nearest neighbor "
    "search, and session management with user accounts and private document collections."
)

add_heading_custom("7.4 Enhanced File Format Support", level=2, italic=True)
add_para(
    "Future versions could support Microsoft Word (.docx) documents, Rich Text Format (.rtf), "
    "Markdown (.md) files, and Optical Character Recognition (OCR) for scanned PDFs."
)

add_heading_custom("7.5 Advanced IR Techniques", level=2, italic=True)
add_para(
    "BM25, a probabilistic retrieval framework that often outperforms TF-IDF, could be "
    "implemented alongside query expansion through pseudo-relevance feedback and result "
    "diversification to ensure result sets cover multiple aspects of a query."
)

# ==================== 8. CONCLUSION ====================
add_heading_custom("8. Conclusion", level=1, center=True, size=11)
add_para(
    "Smart Document Search demonstrates that effective information retrieval systems can be "
    "built with modern web technologies without requiring heavy infrastructure. The TF-IDF "
    "algorithm, despite its age, provides robust and interpretable ranking for document search "
    "tasks. The combination of React's component model, Vite's fast development experience, "
    "and Express's simplicity creates a productive full-stack development environment."
)
add_para(
    "The application's modular architecture and AI SDK integration position it well for "
    "evolution into a semantic search platform. The glassmorphism UI with motion design "
    "creates an engaging user experience that distinguishes it from traditional enterprise "
    "search interfaces."
)
add_para(
    "This project serves as a foundation for understanding information retrieval concepts "
    "while providing a practical tool for personal document management and search."
)

# ==================== REFERENCES ====================
add_heading_custom("References", level=1, center=True, size=11)

references = [
    "[1] Manning, C. D., Raghavan, P., & Schutze, H. (2008). Introduction to Information Retrieval. Cambridge University Press.",
    "[2] Salton, G., Wong, A., & Yang, C. S. (1975). A Vector Space Model for Automatic Indexing. Communications of the ACM, 18(11), 613-620.",
    "[3] Sparck Jones, K. (1972). A Statistical Interpretation of Term Specificity and Its Application in Retrieval. Journal of Documentation, 28(1), 11-21.",
    "[4] Fielding, R. T. (2000). Architectural Styles and the Design of Network-based Software Architectures. Doctoral dissertation, University of California, Irvine.",
    "[5] React Documentation. (2024). React \u2013 The library for web and native user interfaces. Meta Platforms, Inc. https://react.dev",
    "[6] Mozilla Foundation. (2024). PDF.js. https://github.com/mozilla/pdf.js",
    "[7] Google. (2024). Gemini API Documentation. Google AI. https://ai.google.dev/",
    "[8] Vite. (2024). Next Generation Frontend Tooling. https://vitejs.dev/",
    "[9] Tailwind CSS. (2024). Rapidly build modern websites without ever leaving your HTML. https://tailwindcss.com/",
]

for ref in references:
    add_para(ref, size=9, space_after=3)

# ==================== SAVE ====================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "RESEARCH_PAPER.docx")
doc.save(output_path)
print(f"Research paper saved to: {output_path}")